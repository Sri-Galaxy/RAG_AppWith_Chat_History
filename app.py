import os
import streamlit as st
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.runnables import RunnablePassthrough
import tempfile
import uuid

# For document processing
from pdfminer.high_level import extract_text as extract_text_from_pdf
import docx

# Load environment variables
load_dotenv()

# Set page configuration
st.set_page_config(page_title="RAG Chatbot with History", layout="wide")

# Initialize session state for chat history and documents
if "chat_histories" not in st.session_state:
    st.session_state.chat_histories = {}

if "documents" not in st.session_state:
    st.session_state.documents = []

if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Function to get chat history for a session
def get_session_history(session_id: str) -> BaseChatMessageHistory:
    if session_id not in st.session_state.chat_histories:
        st.session_state.chat_histories[session_id] = ChatMessageHistory()
    return st.session_state.chat_histories[session_id]

# Initialize LLM
@st.cache_resource
def initialize_llm():
    gemini_api_key = os.getenv("GOOGLE_API_KEY")
    if not gemini_api_key:
        st.error("Google API Key not found. Please set it in your .env file.")
        st.stop()
    
    os.environ["GOOGLE_API_KEY"] = gemini_api_key
    return ChatGoogleGenerativeAI(model="gemini-2.5-flash", google_api_key=gemini_api_key)

# Initialize embeddings
@st.cache_resource
def initialize_embeddings():
    # Custom embedding function to handle potential dictionary inputs
    class SafeHuggingFaceEmbeddings(HuggingFaceEmbeddings):
        def embed_query(self, text):
            # Handle case where text might be a dictionary
            if isinstance(text, dict) and 'question' in text:
                text = text['question']
            return super().embed_query(text)
    
    return SafeHuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Function to extract text from PDF
def extract_text_from_pdf_file(file_path):
    try:
        return extract_text_from_pdf(file_path)
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return ""

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

# Function to process uploaded documents
def process_uploaded_documents(uploaded_files):
    documents = []
    for file in uploaded_files:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as temp_file:
            temp_file.write(file.getvalue())
            temp_path = temp_file.name
        
        # Extract text content based on file type
        try:
            file_extension = file.name.split('.')[-1].lower()
            content = ""
            
            if file_extension == 'pdf':
                content = extract_text_from_pdf_file(temp_path)
            elif file_extension == 'docx':
                content = extract_text_from_docx(temp_path)
            elif file_extension == 'txt':
                with open(temp_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                st.warning(f"Unsupported file type: {file_extension}")
                continue
            
            if content.strip():
                documents.append(
                    Document(
                        page_content=content,
                        metadata={"source": file.name}
                    )
                )
                st.success(f"Successfully processed {file.name}")
            else:
                st.warning(f"No content extracted from {file.name}")
        except Exception as e:
            st.error(f"Error processing {file.name}: {str(e)}")
        finally:
            # Clean up the temporary file
            os.unlink(temp_path)
    
    return documents

# Create RAG chain
def create_rag_chain(documents, llm, embeddings):
    try:
        # Create vector store
        vector_store = Chroma.from_documents(documents, embedding=embeddings)
        
        # Create retriever
        retriever = vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 2}  # Retrieve top 2 most relevant documents
        )
    except Exception as e:
        st.error(f"Error creating vector store or retriever: {str(e)}")
        raise e
    
    # Create prompt template
    message = """
    Answer the question using ONLY the provided context. If the answer is not in the context, 
    say "I don't have enough information to answer this question based on the provided documents."
    
    Question: {question}
    
    Context:
    {context}
    """
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a helpful assistant that answers questions based only on the provided context."),
        MessagesPlaceholder(variable_name="history"),
        ("human", message)
    ])
    
    # Create RAG chain
    rag_chain = {
        "context": retriever, 
        "question": lambda x: x["question"] if isinstance(x, dict) else x,
        "history": lambda x: x.get("history", []) if isinstance(x, dict) else []
    } | prompt | llm
    
    # Add message history
    with_history = RunnableWithMessageHistory(
        rag_chain,
        get_session_history,
        input_messages_key="history",
        history_messages_key="history"
    )
    
    return with_history

# Main application
def main():
    st.title("📚 RAG Chatbot with Chat History")
    
    # Sidebar for document upload
    with st.sidebar:
        st.header("Document Upload")
        uploaded_files = st.file_uploader("Upload documents", accept_multiple_files=True, type=["txt", "pdf", "docx"])
        
        if uploaded_files:
            if st.button("Process Documents"):
                # Clear existing documents and process new ones
                with st.spinner("Processing documents..."):
                    st.session_state.documents = process_uploaded_documents(uploaded_files)
                st.success(f"Processed {len(st.session_state.documents)} documents")
        
        # Option to clear chat history
        if st.button("Clear Chat History"):
            st.session_state.chat_histories = {}
            st.session_state.session_id = str(uuid.uuid4())
            st.success("Chat history cleared!")
        
        # Display document count
        st.info(f"Documents loaded: {len(st.session_state.documents)}")
        
        # Display document sources
        if st.session_state.documents:
            st.subheader("Loaded Documents")
            for i, doc in enumerate(st.session_state.documents):
                st.write(f"{i+1}. {doc.metadata.get('source', 'Unknown')}")
    
    # Main chat interface
    if not st.session_state.documents:
        st.info("Please upload and process documents to start chatting.")
    else:
        # Initialize LLM and embeddings
        llm = initialize_llm()
        embeddings = initialize_embeddings()
        
        # Create RAG chain
        rag_chain = create_rag_chain(st.session_state.documents, llm, embeddings)
        
        # Chat interface
        st.subheader("Chat with your documents")
        
        # Display chat history
        chat_history = get_session_history(st.session_state.session_id)
        for message in chat_history.messages:
            if isinstance(message, HumanMessage):
                with st.chat_message("user"):
                    st.write(message.content)
            elif isinstance(message, AIMessage):
                with st.chat_message("assistant"):
                    st.write(message.content)
        
        # Chat input
        user_input = st.chat_input("Ask a question about your documents...")
        if user_input:
            # Display user message
            with st.chat_message("user"):
                st.write(user_input)
            
            # Get response from RAG chain
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        config = {"configurable": {"session_id": st.session_state.session_id}}
                        # Debug information
                        print(f"User input type: {type(user_input)}")
                        print(f"User input: {user_input}")
                        # Wrap user_input in a dictionary with 'question' key
                        response = rag_chain.invoke({"question": user_input}, config=config)
                        st.write(response.content)
                    except Exception as e:
                        st.error(f"Error processing question: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())

if __name__ == "__main__":
    main()